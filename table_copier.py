"""
PATCH: table_copier.py
Các bug được fix:
  1. TABLE_PLACEHOLDER_RE hardcode → dynamic từ tables_data
  2. vMerge logic sai → ô phụ merge dọc bị skip hoàn toàn
  3. _create_cell_xml vMerge condition sai (rs==1 không nên có vMerge)
  4. _format_cell_value date dùng Excel format string thay vì strftime format
  5. _parse_hidden_cols chỉ nhận column letter, thêm hỗ trợ tên header
"""

import re
from pathlib import Path
import openpyxl
from docx import Document
from lxml import etree
from config import AppConfig

# FIX 1: Bỏ hardcode — placeholder pattern động, nhận bất kỳ tên nào trong {{...}}
TABLE_PLACEHOLDER_RE = re.compile(r"\{\{\s*(\w+)\s*\}\}")

EXCEL_BORDER_MAP = {
    "thin": "single", "medium": "single", "thick": "single",
    "double": "double", "dashed": "dashed", "dotted": "dotted",
    "hair": "single", None: "none",
}
EXCEL_BORDER_SIZE = {
    "thin": 4, "medium": 8, "thick": 12,
    "double": 6, "dashed": 4, "dotted": 4, "hair": 2,
}


def copy_tables_for_file(
    doc_path: Path,
    config: AppConfig,
    goi_thau_id: str,
    tables_data: list[dict],
    xlsx_path: Path,
    key_id: str = "GoiThau_ID"
):
    doc = Document(str(doc_path))

    file_stem = doc_path.stem
    if file_stem.endswith("-Template"):
        file_stem = file_stem[: -len("-Template")]

    matching_tables = [
        t for t in tables_data
        if str(t.get(key_id) if key_id in t else t.get("GoiThau_ID", "")).strip() == goi_thau_id
        and _match_word(str(t.get("Word", "")), file_stem)
    ]
    if not matching_tables:
        print(f"⚠️  Không tìm thấy bảng nào cho file '{file_stem}' và gói thầu '{goi_thau_id}'")
        return

    # FIX 1: Build set placeholder keys từ tables_data thay vì hardcode
    valid_placeholder_keys = {
        _normalize_placeholder_key(t.get("Name", ""))
        for t in matching_tables
        if t.get("Name")
    }

    doc_occurrences = _collect_table_placeholders(doc, valid_placeholder_keys)

    placeholder_occurrences: dict[str, list] = {}
    for t in matching_tables:
        key = _normalize_placeholder_key(t.get("Name", ""))
        if key not in placeholder_occurrences:
            placeholder_occurrences[key] = []
        placeholder_occurrences[key].append(t)

    for placeholder_key, occurrences in placeholder_occurrences.items():
        doc_occ_list = doc_occurrences.get(placeholder_key, [])
        if not doc_occ_list:
            continue

        if len(occurrences) != len(doc_occ_list):
            print(f"⚠️  '{placeholder_key}': {len(occurrences)} dòng Tables nhưng "
                  f"{len(doc_occ_list)} lần xuất hiện trong Word")

        for i, para_element in enumerate(doc_occ_list):
            if i < len(occurrences):
                occ = occurrences[i]
            else:
                occ = occurrences[-1]

            sheet = str(occ.get("Sheet", "")).strip()
            range_spec = str(occ.get("Range", "A1")).strip()
            hide_cols = str(occ.get("Hide", "")).strip()
            if not sheet:
                continue

            # Tự động phát hiện cột 'File' (chuẩn mới tables-2) và fallback về xlsx_path (chuẩn cũ)
            source_file = xlsx_path
            file_val = str(occ.get("File", "")).strip()
            if file_val and file_val.lower() != "none" and file_val != "":
                resolved_file = xlsx_path.parent / file_val
                if resolved_file.exists():
                    source_file = resolved_file
                else:
                    print(f"⚠️  Không tìm thấy file nguồn cấu hình: {file_val}, sử dụng mặc định: {xlsx_path.name}")

            ws_data = _read_excel_range(source_file, sheet, range_spec, hide_cols)
            if ws_data is None:
                continue

            _insert_table_at_paragraph(doc, para_element, ws_data, config)

    doc.save(str(doc_path))


def _normalize_placeholder_key(raw: str) -> str:
    """Chuẩn hóa key: DanhMuc / {DanhMuc} / {{DanhMuc}} → đều thành 'DanhMuc'"""
    key = raw.strip()
    key = key.strip("{").strip("}")
    return key


def _match_word(tables_word: str, file_stem: str) -> bool:
    """
    FIX: Matching more lenient - partial match instead of exact.
    Handle variations in naming conventions.
    """
    t = tables_word.strip().lower()
    f = file_stem.strip().lower()
    
    # Remove leading numbers: "3. yeu cau" or "9.1 bc" or "18.2. bb" -> "yeu cau" or "bc" or "bb"
    f_no_prefix = re.sub(r"^\d+(\.\d+)*\.?\s*", "", f).strip()
    t_no_prefix = re.sub(r"^\d+(\.\d+)*\.?\s*", "", t).strip()
    
    # Exact match (best)
    if t == f or t == f_no_prefix or t_no_prefix == f or t_no_prefix == f_no_prefix:
        return True
    
    # Partial match - check if one contains the other (handle variations)
    # e.g., "4. BB HDMS.1" matches "4. BB HDMS" or "BB HDMS"
    if len(t) > 3 and len(f) > 3:
        # Remove special chars for comparison
        t_clean = re.sub(r"[^\w\s]", "", t_no_prefix)
        f_clean = re.sub(r"[^\w\s]", "", f_no_prefix)
        
        # Check if one contains the other (at least 5 chars)
        if len(t_clean) > 5 and len(f_clean) > 5:
            if t_clean in f_clean or f_clean in t_clean:
                return True
    
    return False


def _collect_table_placeholders(doc: Document, valid_keys: set) -> dict[str, list]:
    """Collect theo đúng thứ tự xuất hiện trong document"""
    occurrences: dict[str, list] = {}

    def _check_para(para):
        m = TABLE_PLACEHOLDER_RE.search(para.text)
        if m:
            key = m.group(1)
            if key in valid_keys:
                if key not in occurrences:
                    occurrences[key] = []
                occurrences[key].append(para._element)

    for para in doc.paragraphs:
        _check_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _check_para(para)

    return occurrences


def _read_excel_range(xlsx_path: Path, sheet_name: str, range_spec: str, hide_cols: str):
    try:
        wb = openpyxl.load_workbook(xlsx_path, read_only=False, data_only=True)
        # FIX: So sanh ten sheet sau khi strip whitespace (vi du: ' S.DoDa' != 'S.DoDa')
        sheet_name_clean = sheet_name.strip()
        actual_sheet = None
        for s in wb.sheetnames:
            if s.strip() == sheet_name_clean:
                actual_sheet = s
                break
        if actual_sheet is None:
            wb.close()
            return None
        ws = wb[actual_sheet]

        parsed = _parse_range(range_spec, ws)
        if parsed is None:
            wb.close()
            return None
        min_row, max_row, min_col, max_col = parsed

        # FIX 5: _parse_hidden_cols nhận cả column letter VÀ tên header
        # Đọc header row để map tên → index
        header_row = {
            str(ws.cell(row=min_row, column=c).value or "").strip(): c
            for c in range(min_col, max_col + 1)
        }
        hidden_cols = _parse_hidden_cols(hide_cols, min_col, max_col, header_row)
        visible_cols = [c for c in range(min_col, max_col + 1) if c not in hidden_cols]
        col_map = {orig: new for new, orig in enumerate(visible_cols)}

        merged_map = _build_merged_map(ws, min_row, max_row, min_col, max_col)

        # Remap merged cells theo visible columns
        visible_merged: dict[tuple, tuple] = {}
        for (r, c), (mr, mc, rs, cs) in merged_map.items():
            if c in hidden_cols:
                continue
            if mc in hidden_cols:
                # Master bị ẩn → thu gọn span
                new_c = col_map.get(c)
                if new_c is not None:
                    visible_merged[(r - min_row, new_c)] = (r - min_row, new_c, rs, 1)
                continue
            new_r = r - min_row
            new_mr = mr - min_row
            new_c = col_map.get(c)
            new_mc = col_map.get(mc)
            if new_c is not None and new_mc is not None:
                hidden_in_span = len([h for h in hidden_cols if mc <= h < mc + cs])
                new_cs = max(1, cs - hidden_in_span)
                visible_merged[(new_r, new_c)] = (new_mr, new_mc, rs, new_cs)

        row_heights = {}
        for r in range(min_row, max_row + 1):
            h = ws.row_dimensions[r].height
            if h:
                row_heights[r - min_row] = h

        col_widths = {}
        for c in visible_cols:
            w = ws.column_dimensions[openpyxl.utils.get_column_letter(c)].width
            if w:
                col_widths[col_map[c]] = w

        data = {
            "rows": [],
            "merged": visible_merged,
            "row_heights": row_heights,
            "col_widths": col_widths,
        }
        for r in range(min_row, max_row + 1):
            row_data = []
            for c in visible_cols:
                cell = ws.cell(row=r, column=c)
                row_data.append({
                    "value": cell.value,
                    "font": cell.font,
                    "fill": cell.fill,
                    "border": cell.border,
                    "alignment": cell.alignment,
                    "number_format": cell.number_format,
                })
            data["rows"].append(row_data)

        wb.close()
        return data
    except Exception as e:
        print(f"❌ _read_excel_range error: {e}")
        return None


def _parse_range(range_spec: str, ws):
    range_spec = range_spec.strip().upper()
    m = re.match(r"^([A-Z]+)(\d+)(?::([A-Z]+)(\d*))?$", range_spec)
    if not m:
        return None
    min_col = openpyxl.utils.column_index_from_string(m.group(1))
    min_row = int(m.group(2))
    if m.group(3):
        max_col = openpyxl.utils.column_index_from_string(m.group(3))
        max_row = int(m.group(4)) if m.group(4) else ws.max_row
    else:
        max_col = ws.max_column
        max_row = ws.max_row
    return min_row, max_row, min_col, max_col


def _parse_hidden_cols(hide_str: str, min_col: int, max_col: int,
                        header_row: dict[str, int]) -> set[int]:
    """
    FIX 5: Hỗ trợ cả 2 dạng:
    - Column letter: "A,B,C" hoặc "D,E"
    - Tên header:    "DonGia,ThanhTien"
    """
    if not hide_str or hide_str.strip() == "":
        return set()
    cols: set[int] = set()
    for part in hide_str.split(","):
        part = part.strip()
        if not part:
            continue
        # Thử parse column letter trước
        try:
            col_idx = openpyxl.utils.column_index_from_string(part)
            if min_col <= col_idx <= max_col:
                cols.add(col_idx)
            continue
        except Exception:
            pass
        # Thử match tên header (case-insensitive)
        for header_name, col_idx in header_row.items():
            if header_name.lower() == part.lower():
                cols.add(col_idx)
                break
    return cols


def _build_merged_map(ws, min_row, max_row, min_col, max_col) -> dict:
    merged_map = {}
    for merged_range in ws.merged_cells.ranges:
        if (merged_range.min_row > max_row or merged_range.max_row < min_row
                or merged_range.min_col > max_col or merged_range.max_col < min_col):
            continue
        mr = merged_range.min_row
        mc = merged_range.min_col
        rs = merged_range.max_row - merged_range.min_row + 1
        cs = merged_range.max_col - merged_range.min_col + 1
        for r in range(merged_range.min_row, merged_range.max_row + 1):
            for c in range(merged_range.min_col, merged_range.max_col + 1):
                if min_row <= r <= max_row and min_col <= c <= max_col:
                    merged_map[(r, c)] = (mr, mc, rs, cs)
    return merged_map


def _insert_table_at_paragraph(doc: Document, para_element, ws_data, config: AppConfig):
    rows_data = ws_data["rows"]
    if not rows_data:
        return
    nrows = len(rows_data)
    ncols = len(rows_data[0])
    if ncols == 0:
        return

    merged = ws_data.get("merged", {})
    row_heights = ws_data.get("row_heights", {})
    col_widths = ws_data.get("col_widths", {})

    try:
        tbl = _create_word_table_xml(nrows, ncols, rows_data, merged, row_heights, col_widths, config.ExcelToWordWidthFactor)
        parent = para_element.getparent()
        if parent is not None:
            idx = list(parent).index(para_element)
            parent.remove(para_element)
            parent.insert(idx, tbl)
    except Exception as e:
        print(f"⚠️  Lỗi tạo bảng XML: {e}")
        # Keep the placeholder if table creation fails, don't delete it


def _create_word_table_xml(nrows, ncols, rows_data, merged, row_heights, col_widths, factor: int = 90):
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    table_xml = etree.Element(f"{{{NS}}}tbl")

    tblPr = etree.SubElement(table_xml, f"{{{NS}}}tblPr")
    tblStyle = etree.SubElement(tblPr, f"{{{NS}}}tblStyle")
    tblStyle.set(f"{{{NS}}}val", "TableGrid")
    # Tính grid widths trước để biết tổng
    grid_twips = [_col_twips(col_widths, ci, factor=factor) for ci in range(ncols)]
    total_twips = sum(grid_twips)

    tblW = etree.SubElement(tblPr, f"{{{NS}}}tblW")
    tblW.set(f"{{{NS}}}w", str(total_twips))
    tblW.set(f"{{{NS}}}type", "dxa")

    tblGrid = etree.SubElement(table_xml, f"{{{NS}}}tblGrid")

    for ci in range(ncols):
        gridCol = etree.SubElement(tblGrid, f"{{{NS}}}gridCol")
        gridCol.set(f"{{{NS}}}w", str(grid_twips[ci]))

    # FIX 2: Build master cell lookup để biết đâu là master, đâu là phụ
    # merged[(ri, ci)] = (master_ri, master_ci, row_span, col_span)
    master_cells: set[tuple] = set()
    slave_cells: dict[tuple, str] = {}  # (ri,ci) → "h" (horizontal) hoặc "v" (vertical)

    for (ri, ci), (mr, mc, rs, cs) in merged.items():
        if ri == mr and ci == mc:
            master_cells.add((ri, ci))
        elif ri != mr:
            # Ô phụ theo chiều dọc
            slave_cells[(ri, ci)] = "v"
        # Ô phụ theo chiều ngang được Word tự xử lý qua gridSpan → skip

    for ri in range(nrows):
        tr = etree.SubElement(table_xml, f"{{{NS}}}tr")
        trPr = etree.SubElement(tr, f"{{{NS}}}trPr")
        rh = row_heights.get(ri)
        if rh:
            trHeight = etree.SubElement(trPr, f"{{{NS}}}trHeight")
            trHeight.set(f"{{{NS}}}val", str(int(rh * 20)))
            trHeight.set(f"{{{NS}}}hRule", "atLeast")

        for ci in range(ncols):
            merge_info = merged.get((ri, ci))
            is_master = (ri, ci) in master_cells
            slave_type = slave_cells.get((ri, ci))

            # FIX 2: Xác định ô có nằm trong merged range theo chiều ngang không
            # (ô phụ ngang = không phải master nhưng cùng row với master, col > master_col)
            is_h_slave = False
            if merge_info and not is_master and slave_type != "v":
                mr, mc, rs, cs = merge_info
                if ri == mr and ci != mc:
                    is_h_slave = True

            # Bỏ qua ô phụ ngang (Word xử lý tự động qua gridSpan của master)
            if is_h_slave:
                continue

            cell_data = (
                rows_data[ri][ci]
                if ri < len(rows_data) and ci < len(rows_data[ri])
                else None
            )

            # Tính cell width theo col_widths thực tế
            if is_master and merge_info:
                _, _, _, cs = merge_info
                cell_w_twips = sum(
                    _col_twips(col_widths, c, factor=factor)
                    for c in range(ci, ci + cs)
                )
            else:
                cell_w_twips = _col_twips(col_widths, ci, factor=factor)

            if slave_type == "v":
                # FIX 2: Ô phụ merge dọc → tạo ô với vMerge=continue (KHÔNG skip)
                tc = _create_vmerge_continue_cell(NS, cell_w_twips)
                tr.append(tc)
            else:
                # Master cell hoặc ô bình thường
                tc = _create_cell_xml(NS, cell_data, merge_info if is_master else None, cell_w_twips)
                tr.append(tc)

    return table_xml


def _col_twips(col_widths: dict, ci: int, default: int = 1440, factor: int = 90) -> int:
    """Chuyen doi Excel column width -> Word twips.
    Dùng tham số factor cấu hình từ config."""
    cw = col_widths.get(ci)
    if cw and cw > 0:
        return max(200, int(cw * factor))
    return default


def _create_vmerge_continue_cell(NS: str, cell_w_twips: int = 1440):
    """FIX 2: Tạo ô phụ merge dọc với vMerge=continue"""
    tc = etree.Element(f"{{{NS}}}tc")
    tcPr = etree.SubElement(tc, f"{{{NS}}}tcPr")
    tcW = etree.SubElement(tcPr, f"{{{NS}}}tcW")
    tcW.set(f"{{{NS}}}w", str(cell_w_twips))
    tcW.set(f"{{{NS}}}type", "dxa")
    vMerge = etree.SubElement(tcPr, f"{{{NS}}}vMerge")
    vMerge.set(f"{{{NS}}}val", "continue")
    # Thêm paragraph trống bắt buộc
    etree.SubElement(tc, f"{{{NS}}}p")
    return tc


def _create_cell_xml(NS: str, cell_data, merge_info, cell_w_twips: int = 1440):
    tc = etree.Element(f"{{{NS}}}tc")
    tcPr = etree.SubElement(tc, f"{{{NS}}}tcPr")
    tcW = etree.SubElement(tcPr, f"{{{NS}}}tcW")
    tcW.set(f"{{{NS}}}w", str(cell_w_twips))
    tcW.set(f"{{{NS}}}type", "dxa")

    if merge_info:
        mr, mc, rs, cs = merge_info
        # Merge ngang
        if cs > 1:
            gridSpan = etree.SubElement(tcPr, f"{{{NS}}}gridSpan")
            gridSpan.set(f"{{{NS}}}val", str(cs))
        # FIX 3: vMerge chỉ set khi rs > 1 (merge dọc thật sự), và chỉ cho master
        if rs > 1:
            vMerge = etree.SubElement(tcPr, f"{{{NS}}}vMerge")
            vMerge.set(f"{{{NS}}}val", "restart")
        # rs == 1: không có vMerge (ô bình thường, hoặc chỉ merge ngang)

    if cell_data:
        _apply_cell_formatting(tcPr, NS, cell_data)

    p = etree.SubElement(tc, f"{{{NS}}}p")
    pPr = etree.SubElement(p, f"{{{NS}}}pPr")
    if cell_data and cell_data.get("alignment"):
        align = cell_data["alignment"]
        if align.horizontal and align.horizontal != "general":
            jc = etree.SubElement(pPr, f"{{{NS}}}jc")
            jc.set(f"{{{NS}}}val", align.horizontal)

    if cell_data and cell_data.get("value") is not None:
        text_val = _format_cell_value(cell_data)
        r_elem = etree.SubElement(p, f"{{{NS}}}r")
        rPr = etree.SubElement(r_elem, f"{{{NS}}}rPr")

        font = cell_data.get("font")
        if font:
            if font.name:
                rFonts = etree.SubElement(rPr, f"{{{NS}}}rFonts")
                rFonts.set(f"{{{NS}}}ascii", font.name)
                rFonts.set(f"{{{NS}}}hAnsi", font.name)
            if font.size:
                sz = etree.SubElement(rPr, f"{{{NS}}}sz")
                sz.set(f"{{{NS}}}val", str(int(font.size * 2)))
                szCs = etree.SubElement(rPr, f"{{{NS}}}szCs")
                szCs.set(f"{{{NS}}}val", str(int(font.size * 2)))
            if font.bold:
                etree.SubElement(rPr, f"{{{NS}}}b")
            if font.italic:
                etree.SubElement(rPr, f"{{{NS}}}i")
            if font.underline and font.underline not in ("none", False, None):
                u = etree.SubElement(rPr, f"{{{NS}}}u")
                u.set(f"{{{NS}}}val", "single" if font.underline is True else str(font.underline))
            if font.color and hasattr(font.color, "rgb") and font.color.rgb:
                rgb = str(font.color.rgb)
                if rgb and rgb != "00000000" and len(rgb) >= 6:
                    color_elem = etree.SubElement(rPr, f"{{{NS}}}color")
                    color_elem.set(f"{{{NS}}}val", rgb[-6:])  # bỏ alpha prefix

        t = etree.SubElement(r_elem, f"{{{NS}}}t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = str(text_val)

    return tc


def _apply_cell_formatting(tcPr, NS: str, cell_data):
    fill = cell_data.get("fill")
    if fill and fill.fgColor and hasattr(fill.fgColor, "rgb"):
        rgb = str(fill.fgColor.rgb or "")
        # Bỏ qua màu đen mặc định và màu trong suốt
        if rgb and rgb not in ("00000000", "FF000000", "000000", ""):
            shd = etree.SubElement(tcPr, f"{{{NS}}}shd")
            shd.set(f"{{{NS}}}val", "clear")
            shd.set(f"{{{NS}}}color", "auto")
            shd.set(f"{{{NS}}}fill", rgb[-6:])

    tcBorders = etree.SubElement(tcPr, f"{{{NS}}}tcBorders")
    border = cell_data.get("border")
    if border:
        for edge in ("top", "bottom", "left", "right"):
            b = getattr(border, edge, None)
            if b and b.style:
                be = etree.SubElement(tcBorders, f"{{{NS}}}{edge}")
                be.set(f"{{{NS}}}val", EXCEL_BORDER_MAP.get(b.style, "single"))
                be.set(f"{{{NS}}}sz", str(EXCEL_BORDER_SIZE.get(b.style, 4)))
                be.set(f"{{{NS}}}space", "0")
                if b.color and hasattr(b.color, "rgb") and b.color.rgb:
                    be.set(f"{{{NS}}}color", str(b.color.rgb)[-6:])
                else:
                    be.set(f"{{{NS}}}color", "000000")


def _format_cell_value(cell_data) -> str:
    """
    FIX 4: Excel number_format dùng Excel syntax (DD/MM/YYYY),
    cần convert sang Python strftime (%d/%m/%Y) trước khi dùng.
    """
    value = cell_data.get("value")
    if value is None:
        return ""

    nf = str(cell_data.get("number_format") or "")

    if isinstance(value, (int, float)):
        nf_upper = nf.upper()
        # Kiểm tra có phải date format không
        if any(x in nf_upper for x in ("DD", "MM", "YY", "YYYY")):
            from datetime import timedelta, datetime
            base = datetime(1899, 12, 30)
            try:
                dt = base + timedelta(days=float(value))
                # FIX 4: Convert Excel format → Python strftime format
                py_fmt = _excel_date_fmt_to_strftime(nf)
                return dt.strftime(py_fmt)
            except Exception:
                return str(value)

        if "%" in nf:
            return f"{float(value) * 100:.2f}%"

        if "#,##" in nf or "#.##" in nf:
            return f"{float(value):,.0f}".replace(",", ".")

        # Số nguyên không có format đặc biệt
        if value == int(value):
            return str(int(value))
        return str(value)

    return str(value)


def _excel_date_fmt_to_strftime(excel_fmt: str) -> str:
    """
    FIX 4: Convert Excel date format string sang Python strftime.
    Excel: DD/MM/YYYY → Python: %d/%m/%Y
    """
    # Mapping theo thứ tự dài → ngắn để tránh replace nhầm
    mapping = [
        ("YYYY", "%Y"),
        ("YYY",  "%Y"),
        ("YY",   "%y"),
        ("MMMM", "%B"),
        ("MMM",  "%b"),
        ("MM",   "%m"),
        ("M",    "%m"),
        ("DDDD", "%A"),
        ("DDD",  "%a"),
        ("DD",   "%d"),
        ("D",    "%d"),
        ("HH",   "%H"),
        ("H",    "%H"),
        ("SS",   "%S"),
        ("S",    "%S"),
    ]
    result = excel_fmt.upper()
    for excel_token, py_token in mapping:
        result = result.replace(excel_token, py_token)

    # Fallback nếu còn token chưa convert
    if "%" not in result:
        return "%d/%m/%Y"
    return result
