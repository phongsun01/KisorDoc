"""
tests/test_migrator.py
───────────────────────
Unit test cho kisorlib/migrator.py.
"""

import zipfile
from pathlib import Path

import pytest
from docx import Document

from kisorlib.migrator import (
    FileResult,
    _map_placeholder,
    _apply_modifier,
    format_summary,
    migrate_file,
    migrate_folder,
    migrate_xml,
)


# ── helpers ──────────────────────────────────────────────────────────────────

def _make_docx(path: Path, extra_xml: str = "") -> Path:
    doc = Document()
    doc.add_paragraph("placeholder")
    doc.save(str(path))
    if extra_xml:
        tmp = path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(path, "r") as zin:
            with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/document.xml":
                        xml = data.decode("utf-8")
                        xml = xml.replace("</w:body>", f"{extra_xml}</w:body>")
                        data = xml.encode("utf-8")
                    zout.writestr(item, data)
        tmp.replace(path)
    return path


# ══════════════════════════════════════════════
# _apply_modifier
# ══════════════════════════════════════════════

class TestApplyModifier:
    @pytest.mark.parametrize("base,mod,expected", [
        ("NgayKy",       "date",      "{{NgayKy_Date|date}}"),
        ("NgayKy_Date",  "date",      "{{NgayKy_Date|date}}"),   # không double _Date
        ("NgayKy",       "date_long", "{{NgayKy_Date|date_long}}"),
        ("NgayKy",       "day",       "{{NgayKy_Date|day}}"),
        ("NgayKy",       "month",     "{{NgayKy_Date|month}}"),
        ("NgayKy",       "year",      "{{NgayKy_Date|year}}"),
        ("GiaTri",       "number",    "{{GiaTri|number}}"),
        ("GiaTri",       "num2text",  "{{GiaTri|num2text}}"),
        ("TenBien",      "upper",     "{{TenBien|upper}}"),
    ])
    def test_modifier(self, base, mod, expected):
        assert _apply_modifier(base, mod) == expected

    def test_date_not_doubled(self):
        # Nếu base đã có _Date → không thêm nữa
        result = _apply_modifier("NgayKy_Date", "date")
        assert result == "{{NgayKy_Date|date}}"
        assert result.count("_Date") == 1


# ══════════════════════════════════════════════
# _map_placeholder
# ══════════════════════════════════════════════

class TestMapPlaceholder:
    @pytest.mark.parametrize("raw,expected_conv", [
        ("TenGoiThau",          "{{TenGoiThau}}"),
        ("NgayKy.Date",         "{{NgayKy_Date|date}}"),
        ("NgayKy.date",         "{{NgayKy_Date|date}}"),         # case-insensitive
        ("NgayKy.Date.Long",    "{{NgayKy_Date|date_long}}"),
        ("NgayKy.Day",          "{{NgayKy_Date|day}}"),
        ("NgayKy.Month",        "{{NgayKy_Date|month}}"),
        ("NgayKy.Year",         "{{NgayKy_Date|year}}"),
        ("GiaTri.Number",       "{{GiaTri|number}}"),
        ("GiaTri.Chu",          "{{GiaTri|num2text}}"),
        ("GiaTri.Text",         "{{GiaTri|num2text}}"),
        ("TenBien.Upper",       "{{TenBien|upper}}"),
        ("TenBien|upper",       "{{TenBien|upper}}"),            # pipe-filter existing
        ("NgayKy|date",         "{{NgayKy_Date|date}}"),         # normalize _Date
        ("GiaTri|number",       "{{GiaTri|number}}"),
    ])
    def test_conversions(self, raw, expected_conv):
        _, conv = _map_placeholder(raw)
        assert conv == expected_conv, f"Input {raw!r}: expected {expected_conv!r}, got {conv!r}"

    def test_returns_original_clean(self):
        orig, _ = _map_placeholder("NgayKy.Date")
        assert orig == "NgayKy.Date"

    def test_whitespace_stripped(self):
        _, conv = _map_placeholder("  TenGoiThau  ")
        assert conv == "{{TenGoiThau}}"


# ══════════════════════════════════════════════
# migrate_xml
# ══════════════════════════════════════════════

class TestMigrateXml:

    # ── <<...>> HTML-encoded ──

    def test_plain_lt_gt(self):
        xml, changes = migrate_xml("&lt;&lt;TenGoiThau&gt;&gt;")
        assert "{{TenGoiThau}}" in xml
        assert len(changes) == 1

    def test_date_suffix(self):
        xml, changes = migrate_xml("&lt;&lt;NgayKy.Date&gt;&gt;")
        assert "{{NgayKy_Date|date}}" in xml

    def test_date_long_suffix(self):
        xml, changes = migrate_xml("&lt;&lt;NgayKy.Date.Long&gt;&gt;")
        assert "{{NgayKy_Date|date_long}}" in xml

    def test_day_suffix(self):
        xml, changes = migrate_xml("&lt;&lt;NgayKy.Day&gt;&gt;")
        assert "{{NgayKy_Date|day}}" in xml

    def test_month_suffix(self):
        xml, changes = migrate_xml("&lt;&lt;NgayKy.Month&gt;&gt;")
        assert "{{NgayKy_Date|month}}" in xml

    def test_year_suffix(self):
        xml, changes = migrate_xml("&lt;&lt;NgayKy.Year&gt;&gt;")
        assert "{{NgayKy_Date|year}}" in xml

    def test_number_suffix(self):
        xml, changes = migrate_xml("&lt;&lt;GiaTri.Number&gt;&gt;")
        assert "{{GiaTri|number}}" in xml

    def test_chu_suffix(self):
        xml, changes = migrate_xml("&lt;&lt;GiaTri.Chu&gt;&gt;")
        assert "{{GiaTri|num2text}}" in xml

    def test_upper_suffix(self):
        xml, changes = migrate_xml("&lt;&lt;TenBien.Upper&gt;&gt;")
        assert "{{TenBien|upper}}" in xml

    # ── <<...>> raw (không encode) ──

    def test_raw_lt_gt_plain(self):
        xml, changes = migrate_xml("<<TenGoiThau>>")
        assert "{{TenGoiThau}}" in xml

    def test_raw_lt_gt_date(self):
        xml, changes = migrate_xml("<<NgayKy.Date>>")
        assert "{{NgayKy_Date|date}}" in xml

    # ── {{...}} existing — chuẩn hóa ──

    def test_normalize_existing_date(self):
        xml, changes = migrate_xml("{{NgayKy.Date}}")
        assert "{{NgayKy_Date|date}}" in xml

    def test_normalize_existing_number(self):
        xml, changes = migrate_xml("{{GiaTri.Number}}")
        assert "{{GiaTri|number}}" in xml

    def test_existing_already_correct_no_change(self):
        xml, changes = migrate_xml("{{TenGoiThau}}")
        # Không có thay đổi nào
        assert changes == []

    def test_existing_with_correct_filter_no_change(self):
        xml, changes = migrate_xml("{{GiaTri|number}}")
        assert changes == []

    # ── {single} brace → {{double}} ──

    def test_single_brace_table(self):
        xml, changes = migrate_xml("{DanhMuc}")
        assert "{{DanhMuc}}" in xml
        assert len(changes) == 1
        assert changes[0].original == "{DanhMuc}"

    def test_single_brace_multiple(self):
        xml, _ = migrate_xml("{DanhMuc} {KHLCNt}")
        assert "{{DanhMuc}}" in xml
        assert "{{KHLCNt}}" in xml

    def test_single_brace_not_double(self):
        # {{TenBien}} không bị migrate thêm lần nữa
        xml, changes = migrate_xml("{{TenBien}}")
        assert xml == "{{TenBien}}"
        assert changes == []

    # ── Mixed content ──

    def test_mixed_all_types(self):
        xml_in = (
            "&lt;&lt;TenGoiThau&gt;&gt; "
            "&lt;&lt;NgayKy.Date&gt;&gt; "
            "{{GiaTri.Number}} "
            "{DanhMuc}"
        )
        xml_out, changes = migrate_xml(xml_in)
        assert "{{TenGoiThau}}" in xml_out
        assert "{{NgayKy_Date|date}}" in xml_out
        assert "{{GiaTri|number}}" in xml_out
        assert "{{DanhMuc}}" in xml_out
        assert len(changes) >= 3

    # ── No placeholder → không thay đổi ──

    def test_plain_text_unchanged(self):
        xml = "Hello world, no placeholders here."
        out, changes = migrate_xml(xml)
        assert out == xml
        assert changes == []

    def test_empty_string(self):
        out, changes = migrate_xml("")
        assert out == ""
        assert changes == []

    # ── Changes tracking ──

    def test_change_original_recorded(self):
        _, changes = migrate_xml("&lt;&lt;NgayKy.Date&gt;&gt;")
        assert changes[0].original == "<<NgayKy.Date>>"

    def test_change_converted_recorded(self):
        _, changes = migrate_xml("&lt;&lt;NgayKy.Date&gt;&gt;")
        assert changes[0].converted == "{{NgayKy_Date|date}}"


# ══════════════════════════════════════════════
# migrate_file
# ══════════════════════════════════════════════

class TestMigrateFile:

    def test_dry_run_no_file_change(self, tmp_path):
        f = _make_docx(tmp_path / "t.docx", "&lt;&lt;TenGoiThau&gt;&gt;")
        before_mtime = f.stat().st_mtime
        result = migrate_file(f, dry_run=True)
        assert result.success
        assert result.changed
        assert f.stat().st_mtime == before_mtime  # file không bị ghi

    def test_dry_run_no_backup(self, tmp_path):
        f = _make_docx(tmp_path / "t.docx", "&lt;&lt;TenGoiThau&gt;&gt;")
        migrate_file(f, dry_run=True, backup=True)
        bak_dir = tmp_path / "bak"
        assert not bak_dir.exists()

    def test_actual_migrate_creates_backup(self, tmp_path):
        f = _make_docx(tmp_path / "t.docx", "&lt;&lt;TenGoiThau&gt;&gt;")
        result = migrate_file(f, dry_run=False, backup=True)
        assert result.success
        assert result.backed_up_to is not None
        assert result.backed_up_to.exists()

    def test_backup_not_duplicated(self, tmp_path):
        f = _make_docx(tmp_path / "t.docx", "&lt;&lt;TenGoiThau&gt;&gt;")
        migrate_file(f, dry_run=False, backup=True)
        migrate_file(f, dry_run=False, backup=True)  # chạy lại
        bak_dir = tmp_path / "bak"
        bak_files = list(bak_dir.glob("*.bak.docx"))
        assert len(bak_files) == 1  # backup không bị tạo lại

    def test_no_change_file_not_modified(self, tmp_path):
        f = _make_docx(tmp_path / "t.docx")  # không có placeholder cũ
        before = f.stat().st_mtime
        result = migrate_file(f, dry_run=False)
        assert result.success
        assert not result.changed

    def test_missing_file_returns_error(self, tmp_path):
        result = migrate_file(tmp_path / "missing.docx")
        assert not result.success
        assert result.error is not None

    def test_output_is_valid_docx(self, tmp_path):
        f = _make_docx(tmp_path / "t.docx", "&lt;&lt;NgayKy.Date&gt;&gt;")
        migrate_file(f, dry_run=False)
        doc = Document(str(f))
        assert doc is not None

    def test_no_tmp_leftover_on_success(self, tmp_path):
        f = _make_docx(tmp_path / "t.docx", "&lt;&lt;TenGoiThau&gt;&gt;")
        migrate_file(f, dry_run=False)
        tmps = list(tmp_path.glob("*.tmp.docx"))
        assert tmps == []

    def test_changes_count(self, tmp_path):
        f = _make_docx(tmp_path / "t.docx",
                       "&lt;&lt;A&gt;&gt;&lt;&lt;B.Date&gt;&gt;{C}")
        result = migrate_file(f, dry_run=True)
        assert len(result.changes) == 3

    def test_on_progress_called(self, tmp_path):
        events = []
        f = _make_docx(tmp_path / "t.docx", "&lt;&lt;TenGoiThau&gt;&gt;")
        migrate_file(f, dry_run=True, on_progress=events.append)
        assert len(events) >= 1
        assert any("info" in str(e) or "success" in str(e) for e in events)


# ══════════════════════════════════════════════
# migrate_folder
# ══════════════════════════════════════════════

class TestMigrateFolder:

    def test_finds_all_docx(self, tmp_path):
        for i in range(3):
            _make_docx(tmp_path / f"t{i}.docx", "&lt;&lt;A&gt;&gt;")
        results = migrate_folder(tmp_path, dry_run=True)
        assert len(results) == 3

    def test_skips_bak_folder(self, tmp_path):
        _make_docx(tmp_path / "t.docx", "&lt;&lt;A&gt;&gt;")
        bak = tmp_path / "bak"
        bak.mkdir()
        _make_docx(bak / "t.bak.docx", "&lt;&lt;B&gt;&gt;")
        results = migrate_folder(tmp_path, dry_run=True)
        assert len(results) == 1

    def test_skips_bak_docx(self, tmp_path):
        _make_docx(tmp_path / "t.docx", "&lt;&lt;A&gt;&gt;")
        _make_docx(tmp_path / "t.bak.docx", "&lt;&lt;B&gt;&gt;")
        results = migrate_folder(tmp_path, dry_run=True)
        assert len(results) == 1

    def test_recursive_finds_subfolders(self, tmp_path):
        sub = tmp_path / "Opt1"
        sub.mkdir()
        _make_docx(sub / "t.docx", "&lt;&lt;A&gt;&gt;")
        _make_docx(tmp_path / "root.docx", "&lt;&lt;B&gt;&gt;")
        results = migrate_folder(tmp_path, dry_run=True, recursive=True)
        assert len(results) == 2

    def test_non_recursive_ignores_subfolders(self, tmp_path):
        sub = tmp_path / "Opt1"
        sub.mkdir()
        _make_docx(sub / "t.docx", "&lt;&lt;A&gt;&gt;")
        _make_docx(tmp_path / "root.docx", "&lt;&lt;B&gt;&gt;")
        results = migrate_folder(tmp_path, dry_run=True, recursive=False)
        assert len(results) == 1

    def test_empty_folder_returns_empty_list(self, tmp_path):
        results = migrate_folder(tmp_path, dry_run=True)
        assert results == []

    def test_on_progress_total_announced(self, tmp_path):
        for i in range(2):
            _make_docx(tmp_path / f"t{i}.docx", "&lt;&lt;A&gt;&gt;")
        events = []
        migrate_folder(tmp_path, dry_run=True, on_progress=events.append)
        totals = [e.get("total") for e in events if e.get("total")]
        assert 2 in totals

    def test_actual_migrate_all_files(self, tmp_path):
        for i in range(3):
            _make_docx(tmp_path / f"t{i}.docx", "&lt;&lt;A&gt;&gt;")
        results = migrate_folder(tmp_path, dry_run=False, backup=False)
        assert all(r.success for r in results)
        assert all(r.changed for r in results)


# ══════════════════════════════════════════════
# format_summary
# ══════════════════════════════════════════════

class TestFormatSummary:

    def _fr(self, name, success=True, changed=True, changes=None, error=None):
        from kisorlib.migrator import PlaceholderChange
        r = FileResult(path=Path(name), success=success, changed=changed,
                       error=error)
        if changes:
            r.changes = [PlaceholderChange(o, c, "word/document.xml")
                         for o, c in changes]
        return r

    def test_success_changed(self):
        fr = self._fr("t.docx", changes=[("<<A>>", "{{A}}")])
        out = format_summary([fr])
        assert "✅" in out
        assert "t.docx" in out
        assert "<<A>>" in out

    def test_success_no_change(self):
        fr = self._fr("t.docx", changed=False)
        out = format_summary([fr])
        assert "⬜" in out

    def test_error(self):
        fr = self._fr("t.docx", success=False, changed=False, error="FileNotFoundError")
        out = format_summary([fr])
        assert "❌" in out
        assert "FileNotFoundError" in out

    def test_dry_run_icon(self):
        from kisorlib.migrator import PlaceholderChange
        fr = self._fr("t.docx", changes=[("<<A>>", "{{A}}")])
        out = format_summary([fr], dry_run=True)
        assert "🔍" in out

    def test_truncates_long_change_list(self):
        changes = [(f"<<V{i}>>", f"{{{{V{i}}}}}") for i in range(15)]
        fr = self._fr("t.docx", changes=changes)
        out = format_summary([fr])
        assert "và" in out and "khác" in out  # truncation message

    def test_empty_list(self):
        out = format_summary([])
        assert out == "(Không có file nào)"

    def test_multiple_files(self):
        frs = [
            self._fr("a.docx", changes=[("<<A>>", "{{A}}")]),
            self._fr("b.docx", changed=False),
            self._fr("c.docx", success=False, changed=False, error="err"),
        ]
        out = format_summary(frs)
        assert "a.docx" in out
        assert "b.docx" in out
        assert "c.docx" in out
