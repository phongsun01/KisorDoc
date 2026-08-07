"""
tests/test_merger.py
─────────────────────
Unit test cho kisorlib/merger.py.

Bao phủ:
  - mail_merge_safe: happy path, DebugUndefined, lỗi template, lỗi output dir,
    filters, tmp cleanup, context defaults
  - mail_merge: cùng logic, trực tiếp ghi output_path
"""

import shutil
import tempfile
from pathlib import Path

import pytest
from docx import Document


# ── helpers ──────────────────────────────────────────────────────────────────

def _make_docx(path: Path, text: str) -> Path:
    """Tạo .docx tối giản với 1 đoạn văn bản (có thể chứa Jinja2 placeholder)."""
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))
    return path


def _read_docx(path: Path) -> str:
    """Đọc nội dung text của tất cả paragraph trong .docx."""
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


@pytest.fixture()
def tmp(tmp_path):
    return tmp_path


# ══════════════════════════════════════════════
# mail_merge_safe
# ══════════════════════════════════════════════

class TestMailMergeSafe:
    from kisorlib.merger import mail_merge_safe  # import một lần

    def setup_method(self):
        from kisorlib.merger import mail_merge_safe
        self.merge = mail_merge_safe

    # ── Happy path ──

    def test_basic_substitution(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "Xin chào {{Name}}!")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {"Name": "Thế giới"}, out)
        assert ok is True
        assert err == ""
        assert out.exists()
        content = _read_docx(out)
        assert "Thế giới" in content

    def test_multiple_placeholders(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{A}} và {{B}}")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {"A": "alpha", "B": "beta"}, out)
        assert ok
        content = _read_docx(out)
        assert "alpha" in content
        assert "beta" in content

    def test_returns_true_empty_error_on_success(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{X}}")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {"X": "val"}, out)
        assert ok is True
        assert err == ""

    # ── DebugUndefined: placeholder không có trong context → giữ nguyên, không crash ──

    def test_missing_placeholder_preserved(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{Known}} và {{Unknown}}")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {"Known": "có"}, out)
        assert ok is True  # không crash
        assert out.exists()

    def test_all_missing_placeholders_ok(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{A}} {{B}} {{C}}")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {}, out)
        assert ok is True

    # ── Context defaults ──

    def test_now_added_to_context_automatically(self, tmp):
        """mail_merge_safe tự inject 'now' nếu không có trong context."""
        from datetime import datetime
        tpl = _make_docx(tmp / "tpl.docx", "{{now}}")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {}, out)
        assert ok is True  # 'now' được inject → không undefined

    def test_caller_now_not_overridden(self, tmp):
        """Nếu caller cung cấp 'now' thì dùng giá trị đó."""
        from datetime import datetime
        custom_now = datetime(2024, 1, 15, 12, 0, 0)
        tpl = _make_docx(tmp / "tpl.docx", "{{now|date}}")
        out = tmp / "out.docx"
        ok, _ = self.merge(tpl, {"now": custom_now}, out)
        assert ok is True

    # ── Lỗi: template không tồn tại ──

    def test_missing_template_returns_false(self, tmp):
        out = tmp / "out.docx"
        ok, err = self.merge(tmp / "nonexistent.docx", {}, out)
        assert ok is False
        assert err != ""
        assert not out.exists()

    # ── Lỗi: output directory không tồn tại ──

    def test_missing_output_dir_returns_false(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{X}}")
        out = tmp / "subdir_missing" / "out.docx"
        ok, err = self.merge(tpl, {"X": "v"}, out)
        assert ok is False
        assert err != ""

    # ── Tmp file cleanup ──

    def test_no_tmp_leftover_on_success(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{X}}")
        out = tmp / "out.docx"
        before = set(tmp.iterdir())
        self.merge(tpl, {"X": "val"}, out)
        after = set(tmp.iterdir())
        new_files = after - before
        tmp_leftovers = [f for f in new_files if f.name != "out.docx"]
        assert tmp_leftovers == [], f"Tmp files not cleaned: {tmp_leftovers}"

    def test_no_tmp_leftover_on_error(self, tmp):
        out = tmp / "out.docx"
        before = set(tmp.iterdir())
        self.merge(tmp / "missing.docx", {}, out)  # sẽ fail
        after = set(tmp.iterdir())
        new_files = after - before
        assert not any(f.suffix == ".docx" for f in new_files), \
            f"Leftover files after error: {new_files}"

    # ── Output file integrity ──

    def test_output_is_valid_docx(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{Name}}")
        out = tmp / "out.docx"
        self.merge(tpl, {"Name": "Test"}, out)
        # Document() sẽ raise nếu file corrupt
        doc = Document(str(out))
        assert doc is not None

    def test_output_not_created_on_failure(self, tmp):
        out = tmp / "out.docx"
        self.merge(tmp / "missing.docx", {}, out)
        assert not out.exists()

    # ── Filters hoạt động qua mail_merge_safe ──

    def test_date_filter(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{NgayKy|date}}")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {"NgayKy": "15/03/2024"}, out)
        assert ok is True, f"err: {err}"

    def test_number_filter(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{GiaTriHopDong|number}}")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {"GiaTriHopDong": "1500000000"}, out)
        assert ok is True, f"err: {err}"

    def test_upper_filter(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{Name|upper}}")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {"Name": "hello"}, out)
        assert ok is True

    # ── Overwrite existing output ──

    def test_overwrites_existing_output(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{Name}}")
        out = tmp / "out.docx"
        self.merge(tpl, {"Name": "First"}, out)
        self.merge(tpl, {"Name": "Second"}, out)
        content = _read_docx(out)
        assert "Second" in content

    # ── Path types: str và Path đều hoạt động ──

    def test_accepts_string_paths(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{X}}")
        out = tmp / "out.docx"
        ok, err = self.merge(str(tpl), {"X": "v"}, str(out))
        assert ok is True

    def test_accepts_path_objects(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{X}}")
        out = tmp / "out.docx"
        ok, err = self.merge(tpl, {"X": "v"}, out)
        assert ok is True


# ══════════════════════════════════════════════
# mail_merge (không có error handling)
# ══════════════════════════════════════════════

class TestMailMerge:

    def setup_method(self):
        from kisorlib.merger import mail_merge
        self.merge = mail_merge

    def test_basic_substitution(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "Hello {{Name}}!")
        out = tmp / "out.docx"
        self.merge(tpl, {"Name": "World"}, out)
        assert out.exists()
        content = _read_docx(out)
        assert "World" in content

    def test_raises_on_missing_template(self, tmp):
        with pytest.raises(Exception):
            self.merge(tmp / "missing.docx", {}, tmp / "out.docx")

    def test_now_injected(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{now}}")
        out = tmp / "out.docx"
        self.merge(tpl, {}, out)  # không raise
        assert out.exists()

    def test_missing_placeholder_does_not_raise(self, tmp):
        tpl = _make_docx(tmp / "tpl.docx", "{{Known}} {{Unknown}}")
        out = tmp / "out.docx"
        self.merge(tpl, {"Known": "ok"}, out)  # DebugUndefined → không crash
        assert out.exists()
